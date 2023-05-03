from thehive4py.api import TheHiveApi
import docx
import datetime
import json
import re
from docx.shared import Cm
import os
from cortexutils.responder import Responder

def Exportdocx(Responder):

    api = TheHiveApi('http://localhost:9000', 'HereYouPutTheApiKey')

    # # CHOOSE THE CASE BY HAND, YOU CAN JUST "caseId" a case and the document will create # #
    case_id = ["thehive:case"]

    ## REQUEST THE CASE ##
    case_response = api.get_case(case_id)
    case_json = case_response.json()


    # PRINT ANY _json VARIABLE  WITH THIS FORMAT TO MAKE IT READABLE
    #print(json.dumps(case_json, indent=4, sort_keys=True))

    ## REQUEST THE CASE OBSERVABLES ##
    case_observables = api.get_case_observables(
        case_json["_id"], query={}, sort=["-startDate"], range="all"
    )
    observables_json = case_observables.json()

    ## REQUEST EVERY CASE TASK ##
    case_tasks = api.get_case_tasks(case_json["_id"])
    tasks_json = case_tasks.json()
    ## REQUEST EVERY TASK LOG OF THE LAST TASK CASE BASED ON 'id' FIELD ##
    task_log = api.get_task_logs(
        tasks_json[0]['id'], query={}, sort=["-startDate"], range="all"
        )
    task_logs_json = task_log.json()

    ##################### DOCX #########################
    # LOAD THE DOCX DOCUMENT TO USE AS A TEMPLATE, THIS DOCUMENT HAS "HEADING" 1 TO 9 DEFINED, HEADER AND FOOTER AS WELL
    doc = docx.Document('plantilla.docx')
    # LOAD THE STYLE FOR THE TABLES
    tableStyle = doc.styles["Table Grid Light"]
    
    # DOCUMENT TITLE
    doc.add_heading("Reporte de caso n° " + str(case_json["caseId"]), level=0)
    doc.add_paragraph("")

    # CASE INFO TABLE
    tabla_caso = doc.add_table(rows=2, cols=4)
    tabla_caso.style = tableStyle
    celda_0_0_caso = tabla_caso.cell(0, 0)
    celda_0_0_caso.text = "Turno:"
    celda_0_1_caso = tabla_caso.cell(0, 1)
    celda_0_1_caso.text = case_json["customFields"]["turno"]["string"]
    celda_0_2_caso = tabla_caso.cell(0, 2)
    celda_0_2_caso.text = "Fecha:"
    celda_0_3_caso = tabla_caso.cell(0, 3)
    celda_0_3_caso.text = datetime.datetime.fromtimestamp(case_json["createdAt"] / 1000).strftime("%Y-%m-%d %H:%M:%S")
    celda_1_0_caso = tabla_caso.cell(1, 0)
    celda_1_0_caso.text = "Caso:"
    celda_1_1_caso = tabla_caso.cell(1, 1)
    celda_1_1_caso.text = str(case_json["caseId"])
    celda_1_2_caso = tabla_caso.cell(1, 2)
    celda_1_2_caso.text = "Analista:"
    celda_1_3_caso = tabla_caso.cell(1, 3)
    celda_1_3_caso.text = case_json["createdBy"].split("@")[0].replace(".", " ").title()

    #APPLY BOLD FOR COLUMN 0 AND 2
    for row in tabla_caso.rows:
        celdas = row.cells[0]
        for p in celdas.paragraphs:
            p.runs[0].bold = True
    for row in tabla_caso.rows:
        celdas = row.cells[2]
        for p in celdas.paragraphs:
            p.runs[0].bold = True

    #####################################################3
    # EVENT INFORMATION TITLE
    doc.add_paragraph("")
    doc.add_heading("Información del evento", level=1)
    doc.add_paragraph("")

    # EVENT INFORMATION TABLE
    tabla_info = doc.add_table(rows=8, cols=2)
    tabla_info.style = tableStyle
    celda_0_0_info = tabla_info.cell(0, 0)
    celda_0_0_info.text = "Nombre de la regla"
    celda_0_2_info = tabla_info.cell(0, 1)
    celda_0_2_info.text = case_json["title"].replace(".", "[.]").replace("http://", "hxxp://")
    celda_1_0_info = tabla_info.cell(1, 0)
    celda_1_0_info.text = "Signature ID"
    celda_1_1_info = tabla_info.cell(1, 1)
    celda_1_1_info.text = case_json["customFields"]["signature-id"]["string"].replace("null", "N/A").replace("None", "N/A")
    celda_1_2_info = tabla_info.cell(1, 2)
    celda_1_2_info.text = "Severidad"

    # SEVERITY MAP
    severidad_mapa = {
        0: ("Bajo", "Heading 6"),
        1: ("Medio", "Heading 7"),
        2: ("Alto", "Heading 8"), 
        3: ("Critíco", "Heading 9"),
    }
    nivel_severidad, estilo = severidad_mapa.get(case_json["severity"], ("Desconocido", "Normal"))
    celda_1_3_info = tabla_info.cell(1, 3)
    celda_1_3_info.text = nivel_severidad
    celda_1_3_info.paragraphs[0].style = estilo

    celda_3_0_info = tabla_info.cell(3, 0)
    celda_3_0_info.text = "Categoría"
    celda_3_1_info = tabla_info.cell(3, 1)
    celda_3_1_info.text = case_json["customFields"]["categoria"]["string"].replace("null", "N/A").replace("None", "N/A")
    celda_4_0_info = tabla_info.cell(4, 0)
    celda_4_0_info.text = "CVE"
    celda_4_1_info = tabla_info.cell(4, 1)
    celda_4_1_info.text = str(case_json["customFields"]["cve"]["string"]).replace("null", "N/A").replace("None", "N/A")
    celda_5_0_info = tabla_info.cell(5, 0)
    celda_5_0_info.text = "Descripción"
    celda_5_1_info = tabla_info.cell(5, 1)
    celda_5_1_info.text = case_json["description"]

    celda_6_0_info = tabla_info.cell(6, 0)
    celda_6_0_info.text = "Evaluación alerta"

    # SUMMARY MAP
    summary_mapa = {
        "TruePositive": ("Positivo", "Heading 9"),
        "FalsePositive": ("Falso Positivo", "Heading 7"),
        "Indeterminate": ("Indeterminado", "Heading 5"), 
        "Other": ("Otro", "Heading 5"),
    }
    estadoResolucion, estiloResolucion = summary_mapa.get(case_json["resolutionStatus"], ("", "Normal"))
    celda_6_1_info = tabla_info.cell(6, 1)
    celda_6_1_info.text = estadoResolucion
    celda_6_1_info.paragraphs[0].style = estiloResolucion
    celda_7_0_info = tabla_info.cell(7, 0)
    celda_7_0_info.text = "Resumen alerta"
    celda_7_1_info = tabla_info.cell(7, 1)
    # IF SUMMARY IS None THEN 
    if case_json["summary"] == None:
        celda_7_1_info.text = "Caso no cerrado"
    else:
        celda_7_1_info.text = case_json["summary"]

    # APPLY BOLD TO COLUMN 0
    for row in tabla_info.rows:
        celdas = row.cells[0]
        for p in celdas.paragraphs:
            p.runs[0].bold = True

    #########################################################

    # OBSERVABLES TITLE ONLY IF THERE ARE ANY
    if observables_json != []:
        doc.add_paragraph("")
        doc.add_heading("Observables del caso", level=1)
        doc.add_paragraph("")

    # FOR EVERY OBSERVABLE ADD 3 ROWS FOR DATA, DATATYPE AND ADDED BY
    for observable in observables_json:
        tabla_obs = doc.add_table(rows=0, cols=2)
        tabla_obs.style = tableStyle
        creador = observable["createdBy"].split("@")[0].replace(".", " ").title()
        artifact = observable["data"].replace(".", "[.]").replace("http://", "http[:]//")
        tipo = observable["dataType"]

        celda_0_0 = tabla_obs.add_row().cells[0]
        celda_0_0.text = "Elemento observable"
        celda_0_0.paragraphs[0].style = "Heading 5"

        celda_0_1 = tabla_obs._cells[1]
        celda_0_1.text = artifact

        celda_1_0 = tabla_obs.add_row().cells[0]
        celda_1_0.text = "Tipo"
        celda_1_0.paragraphs[0].style = "Heading 5"

        celda_1_1 = tabla_obs._cells[3]
        celda_1_1.text = tipo

        celda_2_0 = tabla_obs.add_row().cells[0]
        celda_2_0.text = "Añadido por"
        celda_2_0.paragraphs[0].style = "Heading 5"

        celda_2_1 = tabla_obs._cells[5]
        celda_2_1.text = creador
        i = 7
        # ADD A ROW FOR EVERY ANALYZER AND DISPLAY THE OUTPUT OF MINIREPORT
        for report in observable["reports"]:
            #variables
            report_level = observable["reports"][report]["taxonomies"][0]["level"]
            report_value = observable["reports"][report]["taxonomies"][0]["value"]
            report_predicate = observable["reports"][report]["taxonomies"][0]["predicate"]
            #crea la tabla 
            analyzer = tabla_obs.add_row().cells[0]
            analyzer.text = report.replace("_", "").replace("0", "").replace("1", "").replace("2", "").replace("3", "").replace("4", "")
            analyzer.paragraphs[0].style = "Heading 5"

            analyzer_report = tabla_obs._cells[i]
            analyzer_report.text = "Level: " + str(report_level) + "\nResultado: " + str(report_value) + "\nCriterio: " + str(report_predicate)
            i += 2
        # ADD A SPACE BETWEEN EVERY OBSERVABLE TABLE
        doc.add_paragraph("")

    ####################################################
    # TASK LOGS TITLE
    doc.add_paragraph("")
    doc.add_heading("Registros del caso", level=1)
    doc.add_paragraph("")

    # TASK LOGS TABLE
    tabla_tasklog = doc.add_table(rows=0, cols=1)
    tabla_tasklog.style = tableStyle

    for tarea in task_logs_json:
        # VARIABLE IN TASKS
        fecha = datetime.datetime.fromtimestamp(tarea["createdAt"] / 1000).strftime("%Y-%m-%d %H:%M:%S")
        creador = tarea["createdBy"].split("@")[0].replace(".", " ").title()
        mensaje = tarea["message"]
        # ADD 2 ROW FOR EVERY TASK
        celda1 = tabla_tasklog.add_row().cells[0]
        celda2 = tabla_tasklog.add_row().cells[0]
        celda1.text = f"{fecha} - {creador}"
        celda1.paragraphs[0].style = "Heading 5"
        # IF THERES ANY ATTACHMENT WITH "contentType": "image/jpeg" OR "contentType": "image/png", DOWNLOAD, ADD IT TO DOCX AND DELETE IT
        if tarea.get("attachment") and tarea["attachment"].get("contentType") in ["image/jpeg", "image/png"]:
            
            imagen_id = tarea["attachment"]["id"]
            imagen_query = api.download_attachment(imagen_id)
            imagen_bytes = imagen_query.content
            imagen_archivo = tarea["attachment"]["name"]
            with open(imagen_archivo, "wb") as f:
                f.write(imagen_bytes)
            
            celda2.text = mensaje
            celda2.paragraphs[0].style = "Heading 4"
            celda2.add_paragraph().add_run().add_picture(imagen_archivo, width=Cm(15))
            celda2.add_paragraph().add_run().add_text("\n \n")

            os.remove(imagen_archivo)
        else:
            celda2.text = mensaje
            celda2.paragraphs[0].style = "Heading 4"
    # SAVE THE DOCUMENT AS "CASEID - CASETITLE" REPLACING SPECIAL CHARACTERS WITH "_"
    regex = re.compile('[^a-zA-Z0-9\s]')
    nombre_archivo = str(case_json["caseId"]) + " - " + regex.sub('_', case_json["title"])
    doc.save(nombre_archivo + ".docx")
if __name__=='__main__' :
    Exportdocx()
